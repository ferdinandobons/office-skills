# SPDX-License-Identifier: MIT
"""Deterministic builder for the COMPLEX synthetic DOCX example template.

Produces ``examples/templates/branddocs_template.docx``: a 100% synthetic
(``BrandDocs Corp``, never proprietary) Word template that stresses the brand-docx
extractor / generator across as many Word component types as can be authored
without Microsoft Word - python-docx for the bulk, and raw lxml for the many
parts python-docx cannot reach (block-level ``w:sdt`` content controls, real
TOC / SEQ complex fields, ``word/numbering.xml`` abstractNum/num, a custom
``w:type="table"`` style, header logo drawing, ``PAGE`` field, footnotes, and a
landscape ``w:sectPr``).

Components authored (each is a real OOXML structure, not a text approximation):

  COVER (multi-slot front matter, all in the cover region before the TOC):
    * a block-level ``w:sdt`` TITLE content control with ``w:alias='Title'`` and a
      realistic synthetic title - the shape ``cover.discover_cover`` keys on as an
      SDT anchor (python-docx cannot author this, so it is lxml);
    * a SUBTITLE / description placeholder paragraph;
    * a DOCUMENT-ID placeholder paragraph ("Document ID: {{doc_id}}");
    * a DATE placeholder paragraph ("{{date}}").
    * a compact executive scorecard table with three synthetic KPI tiles.

  INDEX FRONT MATTER (three real complex fields, each with cached demo entries):
    * an outline Table of Contents  ``TOC \\o "1-3" \\h \\z \\u``;
    * a Table of Tables             ``TOC \\h \\z \\c "Table"``;
    * a Table of Figures            ``TOC \\h \\z \\c "Figure"``.
    Each carries a cached result (styled entry paragraphs with PAGEREF) so the
    field renders before Word recomputes it, and each ``\\c`` switch is the
    opaque seq id ``structure.inventory_fields`` surfaces.

  NUMBERING (real ``word/numbering.xml``, referenced via ``w:numPr`` from named
    paragraph styles, not direct formatting):
    * a 2-level BULLET list  -> styles "BrandDocs Bullet L1" / "BrandDocs Bullet L2";
    * a 1-level NUMBERED list -> style "BrandDocs Number L1".

  TABLES: a custom ``w:type="table"`` style "BrandDocs Table" (header-row
    shading + row banding via ``w:tblStylePr`` conditional formatting) applied to
    a sample revenue table, an executive scorecard, a risk/readiness matrix, and
    a landscape rollout matrix, with real ``SEQ Table`` captions on body tables.

  FIGURE: two real inline PNG figures, each with a ``SEQ Figure`` caption -
    Figure 1 is the shared text-only BrandDocs wordmark (image1.png) and Figure 2
    is a real rising growth curve (image2.png, a distinct media part).

  CALLOUT: a paragraph style "BrandDocs Callout" with shading + a box border.
  QUOTE: a paragraph style "BrandDocs Quote" with a navy/amber editorial treatment.

  HEADER / FOOTER: the shared BrandDocs wordmark in the default header, and a
    ``PAGE`` field in the default footer.

  SECTIONS: a PORTRAIT first section and a LANDSCAPE second section (distinct
    ``w:sectPr`` page size + orientation).

  FOOTNOTE: one real footnote (``word/footnotes.xml`` + a referencing run).

  DEMO BODY: instruction / lorem-ipsum body content (an "Example heading"
    Heading-1 and following paragraphs) a generation is expected to clear.

The output is byte-reproducible: every id / image byte / part is fixed, no
random or timestamp, so re-running yields an identical file.

Run:
    PYTHONPATH=scripts .venv/bin/python examples/builders/build_branddocs_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Twips
from lxml import etree

from _brandlib import (
    branddocs_curve_png,
    branddocs_mark_png,
    brand_theme_slots,
    freeze_ooxml,
)

OUT = Path(__file__).resolve().parents[1] / "templates" / "branddocs_template.docx"

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"

# Synthetic BrandDocs Corp brand palette (made-up; never proprietary).
BRAND_NAVY = "16213F"
BRAND_TEAL = "2B7CD3"
BRAND_AMBER = "E0742B"
BRAND_LIGHT = "EAF1FF"
BRAND_BAND = "DCE7FF"
WHITE = "FFFFFF"


# ---------------------------------------------------------------------------
# lxml element helpers
# ---------------------------------------------------------------------------
def _w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def _el(tag: str, **attrs) -> etree._Element:
    """Make a ``w:``-namespaced element with ``w:``-namespaced attributes."""
    e = etree.SubElement(etree.Element(_w("_root")), _w(tag))
    e.getparent().remove(e)
    for k, v in attrs.items():
        e.set(_w(k), v)
    return e


def _sub(parent: etree._Element, tag: str, **attrs) -> etree._Element:
    e = etree.SubElement(parent, _w(tag))
    for k, v in attrs.items():
        e.set(_w(k), v)
    return e


def _run(text: str, *, preserve: bool = True, instr: bool = False) -> etree._Element:
    """A ``w:r`` carrying either a ``w:t`` or a ``w:instrText``."""
    r = _el("r")
    leaf = _sub(r, "instrText" if instr else "t")
    leaf.text = text
    if preserve:
        leaf.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _fldchar(kind: str, *, dirty: bool = False) -> etree._Element:
    r = _el("r")
    fc = _sub(r, "fldChar", fldCharType=kind)
    if dirty and kind == "begin":
        fc.set(_w("dirty"), "true")
    return r


def _add_external_hyperlink(doc, paragraph, url: str, text: str) -> None:
    """Append a real ``w:hyperlink`` (external relationship) carrying ``text``.

    Builds the ``r:id``-namespaced attribute directly (the ``_sub`` helper only sets
    ``w:`` attributes). The link run now carries a ``w:rStyle='Hyperlink'`` (DOCX-A6):
    the template registers a real Hyperlink character style (teal ``w:color`` + single
    underline), so the link color is captured as a ``link.color`` provenance fact.
    """
    rid = doc.part.relate_to(url, f"{R}/hyperlink", is_external=True)
    hl = etree.SubElement(paragraph._p, _w("hyperlink"))
    hl.set(f"{{{R}}}id", rid)
    run = _run(text)
    # Stamp the Hyperlink rStyle at the head of the run's w:rPr.
    rpr = etree.Element(_w("rPr"))
    _sub(rpr, "rStyle", val="Hyperlink")
    run.insert(0, rpr)
    hl.append(run)


def _set_pstyle(paragraph, style_id: str) -> None:
    """Stamp ``w:pPr/w:pStyle@w:val`` on a paragraph by STYLE ID directly.

    Bypasses python-docx's ``style=`` name lookup (which is deprecated for ids and
    cannot reference a custom style not registered under its display name), and
    matches how the extractor keys on style ids.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    existing = pPr.find(_w("pStyle"))
    if existing is not None:
        existing.set(_w("val"), style_id)
        return
    pStyle = etree.SubElement(pPr, _w("pStyle"))
    pStyle.set(_w("val"), style_id)
    pPr.insert(0, pStyle)


def _p(doc, text: str = "", style_id: str | None = None):
    """Add a body paragraph, optionally stamping a style id directly."""
    paragraph = doc.add_paragraph(text)
    if style_id:
        _set_pstyle(paragraph, style_id)
    return paragraph


def _brand_runs(paragraph, *, font=None, color=None, size_hp=None):
    """Stamp DIRECT run-level ``w:rPr`` (font + color + size) on EVERY run of a
    paragraph (DOCX-A4 direct-run typography capture).

    The engine's ``capture_fonts`` / ``capture_palette`` read EXPLICIT run formatting
    (``run.font.name`` / ``run.font.size`` / ``run.font.color.rgb``), so the brand's
    real visible typography is captured per role and as the document body defaults
    only when it lives on the runs - not just in named styles. Uses python-docx's
    ``run.font`` API (the same surface the extractor introspects); a run with no text
    is skipped so empty decorative paragraphs never contribute a phantom value.
    """
    from docx.shared import Pt, RGBColor

    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        if font:
            run.font.name = font
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if size_hp:
            run.font.size = Pt(size_hp / 2)


def _brand_table_body(table, *, header_color=WHITE, body_color=BRAND_NAVY):
    """Stamp DIRECT brand run typography on a body table's cells (DOCX-A4).

    Header-row runs get Arial + ``header_color`` (white over the navy header fill);
    every body-row run gets Calibri + ``body_color`` at 11pt. This makes the table's
    visible figures part of the captured BODY typography population (the dominant
    direct font/size/color the engine writes to ``theme.fonts.body`` / ``theme.text``).
    """
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                if ri == 0:
                    _brand_runs(para, font="Arial", color=header_color)
                else:
                    _brand_runs(para, font="Calibri", color=body_color, size_hp=22)


def _pr(doc, text, style_id=None, *, font=None, color=None, size_hp=None):
    """Add a paragraph then stamp brand DIRECT run typography on its runs.

    A thin wrapper over :func:`_p` + :func:`_brand_runs` so body content carries the
    brand's real Arial/Calibri + navy/teal/amber run formatting the engine captures.
    """
    paragraph = _p(doc, text, style_id)
    _brand_runs(paragraph, font=font, color=color, size_hp=size_hp)
    return paragraph


def _add_styled_run(paragraph, text, rstyle, *, color=None, font=None):
    """Append a run carrying a ``w:rStyle`` character-style id + optional direct
    font/color (DOCX-A6 character-style application).

    Used to put a ``BrandDocsEmphasis`` / ``BrandDocsLeadIn`` mark on a body run so
    the artifact catalog's character styles are actually exercised on content.
    """
    from docx.shared import RGBColor

    run = paragraph.add_run(text)
    run.style = rstyle
    if font:
        run.font.name = font
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


# ---------------------------------------------------------------------------
# THEME (DOCX-A1 + DOCX-A2) - rewrite the EXISTING word/theme/theme1.xml part in
# place: a:fontScheme (Arial major / Calibri minor) + a:clrScheme (BrandDocs
# navy/teal/amber/light). Mutating the shipped part (never adding a second theme
# part) mirrors the hardened in-place pattern used for word/numbering.xml.
# ---------------------------------------------------------------------------
def _a(tag: str) -> str:
    return f"{{{A}}}{tag}"


def _theme_part(doc):
    """The package's ``word/theme/theme1.xml`` Part (python-docx ships exactly one).

    Reached via the package part iterator (python-docx exposes no theme accessor),
    so the bytes can be parsed, rewritten, and written back onto the SAME Part - no
    duplicate part, the zip-corruption chokepoint the numbering path also avoids.
    """
    for part in doc.part.package.iter_parts():
        if str(part.partname) == "/word/theme/theme1.xml":
            return part
    return None


def _rewrite_theme_fontscheme(scheme) -> None:
    """Set ``a:majorFont/a:latin@typeface='Arial'`` and ``a:minorFont`` Calibri.

    Rewrites only the ``a:latin`` typeface of each font; the ``a:ea`` / ``a:cs`` /
    script fallbacks the stock theme ships are left intact. This is the OOXML
    equivalent of the docx skill's document-default run font pairing, authored at
    the theme ``a:fontScheme`` level (read by the engine's ``_extract_theme_fonts``).
    """
    faces = {"majorFont": "Arial", "minorFont": "Calibri"}
    for font_tag, face in faces.items():
        font = scheme.find(_a(font_tag))
        if font is None:
            continue
        latin = font.find(_a("latin"))
        if latin is None:
            latin = etree.SubElement(font, _a("latin"))
            font.insert(0, latin)
        latin.set("typeface", face)


def _rewrite_theme_clrscheme(scheme) -> None:
    """Rewrite the ``a:clrScheme`` children to the BrandDocs palette (DOCX-A2).

    Each slot becomes a static ``a:srgbClr val='RRGGBB'`` (deterministic; no sysClr
    so the bytes are constant). The dk1/lt1 (text/background) pair Word writes as
    ``a:sysClr`` is replaced with explicit srgb so the parsed palette carries the
    BrandDocs navy/light directly. Slots not named by the brand map keep a sensible
    branded value so every clrScheme child resolves to a real BrandDocs-family hex.
    """
    slots = brand_theme_slots()  # dk1/lt1/accent1/accent2/accent4/hlink/folHlink
    # Fill the remaining canonical slots so every child carries an explicit srgb.
    full = {
        "dk1": slots["dk1"],
        "lt1": slots["lt1"],
        "dk2": BRAND_TEAL,
        "lt2": BRAND_LIGHT,
        "accent1": slots["accent1"],
        "accent2": slots["accent2"],
        "accent3": BRAND_BAND,
        "accent4": slots["accent4"],
        "accent5": BRAND_TEAL,
        "accent6": BRAND_AMBER,
        "hlink": slots["hlink"],
        "folHlink": slots["folHlink"],
    }
    for slot, hexval in full.items():
        node = scheme.find(_a(slot))
        if node is None:
            node = etree.SubElement(scheme, _a(slot))
        for child in list(node):
            node.remove(child)
        srgb = etree.SubElement(node, _a("srgbClr"))
        srgb.set("val", hexval)


def _mutate_theme(doc) -> None:
    """In-place rewrite of theme1.xml fontScheme + clrScheme (DOCX-A1 + DOCX-A2).

    Parses the EXISTING theme part, rewrites the two schemes, and writes the bytes
    back onto the same Part via ``Part.blob =``. Pure static XML, so the output
    bytes are constant (``freeze_ooxml`` still pins the zip timestamps).
    """
    part = _theme_part(doc)
    if part is None:
        return
    root = etree.fromstring(part.blob)
    elements = root.find(_a("themeElements"))
    if elements is None:
        return
    clr = elements.find(_a("clrScheme"))
    if clr is not None:
        _rewrite_theme_clrscheme(clr)
    fonts = elements.find(_a("fontScheme"))
    if fonts is not None:
        _rewrite_theme_fontscheme(fonts)
    # The shipped theme1.xml is loaded as a plain ``Part`` whose ``blob`` reads
    # ``_blob``; write the rewritten bytes straight back onto the same part (no
    # duplicate part, the zip-corruption chokepoint).
    part._blob = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )


def _mutate_doc_defaults(doc) -> None:
    """Pin the brand body face/size/color as the docDefaults run default (DOCX-A3).

    Adds (or updates) ``docDefaults/rPrDefault/rPr`` with ``w:rFonts@w:ascii/@w:hAnsi
    ='Calibri'``, ``w:sz='22'`` (11pt), and ``w:color='16213F'`` so the document
    default body run is an explicit literal (read by the engine's ``_doc_default_ascii``
    as the major/minor font fallback). Element order inside ``w:rPr`` is rFonts,
    color, sz per CT_RPr.
    """
    styles = doc.styles.element
    dd = styles.find(_w("docDefaults"))
    if dd is None:
        dd = etree.SubElement(styles, _w("docDefaults"))
        styles.insert(0, dd)
    rpr_default = dd.find(_w("rPrDefault"))
    if rpr_default is None:
        rpr_default = _sub(dd, "rPrDefault")
    rpr = rpr_default.find(_w("rPr"))
    if rpr is None:
        rpr = _sub(rpr_default, "rPr")
    for tag in ("rFonts", "color", "sz"):
        existing = rpr.find(_w(tag))
        if existing is not None:
            rpr.remove(existing)
    rfonts = _sub(rpr, "rFonts")
    rfonts.set(_w("ascii"), "Calibri")
    rfonts.set(_w("hAnsi"), "Calibri")
    _sub(rpr, "color", val=BRAND_NAVY)
    _sub(rpr, "sz", val="22")


# ---------------------------------------------------------------------------
# Custom STYLES (paragraph styles + the branded table style). Authored straight
# into ``word/styles.xml`` via lxml so we control header shading / banding /
# borders python-docx cannot express.
# ---------------------------------------------------------------------------
def _add_paragraph_style(
    styles,
    style_id,
    name,
    *,
    based_on="Normal",
    color=None,
    bold=False,
    size_pt=None,
    shading=None,
    box_border=None,
    left_accent=None,
):
    st = _sub(styles, "style", type="paragraph", styleId=style_id)
    st.set(_w("customStyle"), "1")
    _sub(st, "name", val=name)
    _sub(st, "basedOn", val=based_on)
    _sub(st, "qFormat")
    pPr = _sub(st, "pPr")
    if shading:
        _sub(pPr, "shd", val="clear", color="auto", fill=shading)
    if box_border:
        pbdr = _sub(pPr, "pBdr")
        for side in ("top", "left", "bottom", "right"):
            sz = "12"
            col = box_border
            # BP-CALLOUT-ACCENT: thick amber accent bar on the left edge.
            if side == "left" and left_accent:
                sz = "24"
                col = left_accent
            _sub(pbdr, side, val="single", sz=sz, space="6", color=col)
        _sub(pPr, "spacing", before="120", after="120")
    rPr = _sub(st, "rPr")
    if bold:
        _sub(rPr, "b")
    if color:
        _sub(rPr, "color", val=color)
    if size_pt:
        _sub(rPr, "sz", val=str(int(size_pt * 2)))
    return st


def _add_list_style(styles, style_id, name, num_id, ilvl=0):
    """A list paragraph style that references a w:num via w:numPr.

    ``ilvl`` pins the list level on the style's ``w:numPr`` (``w:ilvl`` before
    ``w:numId``, OOXML order). A second-level bullet style MUST declare
    ``w:ilvl=1`` so ``structure.style_num_binding`` reads level 1 and the role
    resolves to ``list.bullet.2`` (1-based = ilvl+1) instead of colliding with
    the level-1 bullet role on a missing-ilvl default of 0.
    """
    st = _sub(styles, "style", type="paragraph", styleId=style_id)
    st.set(_w("customStyle"), "1")
    _sub(st, "name", val=name)
    _sub(st, "basedOn", val="ListParagraph")
    _sub(st, "qFormat")
    pPr = _sub(st, "pPr")
    numPr = _sub(pPr, "numPr")
    _sub(numPr, "ilvl", val=str(ilvl))
    _sub(numPr, "numId", val=str(num_id))
    return st


def _add_table_style(styles):
    """A custom ``w:type='table'`` style: header-row shading + row banding."""
    st = _sub(styles, "style", type="table", styleId="BrandDocsTable")
    st.set(_w("customStyle"), "1")
    _sub(st, "name", val="BrandDocs Table")
    _sub(st, "basedOn", val="TableNormal")
    _sub(st, "uiPriority", val="99")
    # Base table: thin navy grid + banded-row size hint.
    tblPr = _sub(st, "tblPr")
    _sub(tblPr, "tblStyleRowBandSize", val="1")
    borders = _sub(tblPr, "tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _sub(borders, side, val="single", sz="4", space="0", color=BRAND_NAVY)
    # Default cell run color.
    rPr = _sub(st, "rPr")
    _sub(rPr, "color", val=BRAND_NAVY)
    # First-row (header) conditional formatting: navy fill + white bold text.
    fr = _sub(st, "tblStylePr", type="firstRow")
    fr_rpr = _sub(fr, "rPr")
    _sub(fr_rpr, "b")
    _sub(fr_rpr, "color", val=WHITE)
    fr_tcpr = _sub(fr, "tcPr")
    _sub(fr_tcpr, "shd", val="clear", color="auto", fill=BRAND_NAVY)
    # Banded rows: light-blue fill on every other row.
    band = _sub(st, "tblStylePr", type="band1Horz")
    band_tcpr = _sub(band, "tcPr")
    _sub(band_tcpr, "shd", val="clear", color="auto", fill=BRAND_BAND)
    return st


def _override_builtin_headings(styles):
    """Override the BUILTIN Heading1/2/3 styles by exact styleId (DOCX-A5).

    Restyles them to the brand look (navy bold Arial, brand spacing, keepNext,
    outlineLevel 0/1/2) - the docx skill's named-style override pattern. If a
    builtin element already exists (python-docx emits Heading1/2 lazily once used),
    its run/paragraph props are rewritten in place; otherwise a fresh
    ``w:style styleId='HeadingN'`` is authored with the canonical name / basedOn /
    quickFormat. outlineLevel is REQUIRED for TOC inclusion.
    """
    existing = {st.get(_w("styleId")): st for st in styles.findall(_w("style"))}
    for lvl in (1, 2, 3):
        sid = f"Heading{lvl}"
        size_pt = {1: 18, 2: 15, 3: 13}[lvl]
        color = BRAND_NAVY if lvl < 3 else BRAND_TEAL
        st = existing.get(sid)
        if st is None:
            st = _sub(styles, "style", type="paragraph", styleId=sid)
            _sub(st, "name", val=f"heading {lvl}")
            _sub(st, "basedOn", val="Normal")
            _sub(st, "next", val="Normal")
            _sub(st, "uiPriority", val="9")
            _sub(st, "qFormat")
        # Drop any prior pPr/rPr we are about to re-author (idempotent rewrite).
        for tag in ("pPr", "rPr"):
            old = st.find(_w(tag))
            if old is not None:
                st.remove(old)
        pPr = _sub(st, "pPr")
        _sub(pPr, "keepNext")
        _sub(pPr, "spacing", before="240", after="120")
        _sub(pPr, "outlineLvl", val=str(lvl - 1))
        rPr = _sub(st, "rPr")
        rfonts = _sub(rPr, "rFonts")
        rfonts.set(_w("ascii"), "Arial")
        rfonts.set(_w("hAnsi"), "Arial")
        _sub(rPr, "b")
        _sub(rPr, "color", val=color)
        _sub(rPr, "sz", val=str(size_pt * 2))


def _add_section_number_style(styles):
    """A heading-number variant paragraph style (DOCX-A5): BrandDocsSectionNumber.

    A navy bold Arial run style for a numbered section label that rounds out the
    editorial heading system without colliding with the builtin Heading roles.
    """
    _add_paragraph_style(
        styles,
        "BrandDocsSectionNumber",
        "BrandDocs Section Number",
        color=BRAND_NAVY,
        bold=True,
        size_pt=11,
    )


def _add_character_style(
    styles, style_id, name, *, color=None, bold=False, small_caps=False, underline=False
):
    """Author a NAMED CHARACTER style (``w:type='character'``) via lxml (DOCX-A6).

    Element order inside ``w:rPr`` follows CT_RPr: rStyle is on the run, here we
    emit ``b`` / ``smallCaps`` / ``color`` / ``u`` in schema order (``w:color``
    precedes ``w:u``). Used for the brand emphasis / lead-in marks and the real
    Hyperlink / FollowedHyperlink link-color styles.
    """
    st = _sub(styles, "style", type="character", styleId=style_id)
    st.set(_w("customStyle"), "1")
    _sub(st, "name", val=name)
    _sub(st, "uiPriority", val="1")
    _sub(st, "qFormat")
    rPr = _sub(st, "rPr")
    if bold:
        _sub(rPr, "b")
    if small_caps:
        _sub(rPr, "smallCaps")
    if color:
        _sub(rPr, "color", val=color)
    if underline:
        _sub(rPr, "u", val="single")
    return st


def _add_character_styles(styles):
    """Register the brand character styles + Hyperlink / FollowedHyperlink (DOCX-A6).

    The Hyperlink style is NOT marked ``w:customStyle`` (it is the well-known builtin
    style id the engine's ``_is_link_run`` keys on); its ``w:color`` precedes ``w:u``
    at the schema-correct CT_RPr position so the teal link color is captured as a
    ``link.color`` provenance fact.
    """
    have = {st.get(_w("styleId")) for st in styles.findall(_w("style"))}
    if "BrandDocsEmphasis" not in have:
        _add_character_style(
            styles,
            "BrandDocsEmphasis",
            "BrandDocs Emphasis",
            color=BRAND_TEAL,
            bold=True,
        )
    if "BrandDocsLeadIn" not in have:
        _add_character_style(
            styles,
            "BrandDocsLeadIn",
            "BrandDocs Lead-in",
            color=BRAND_NAVY,
            bold=True,
            small_caps=True,
        )
    if "Hyperlink" not in have:
        st = _sub(styles, "style", type="character", styleId="Hyperlink")
        _sub(st, "name", val="Hyperlink")
        _sub(st, "uiPriority", val="99")
        rPr = _sub(st, "rPr")
        _sub(rPr, "color", val=BRAND_TEAL)
        _sub(rPr, "u", val="single")
    if "FollowedHyperlink" not in have:
        st = _sub(styles, "style", type="character", styleId="FollowedHyperlink")
        _sub(st, "name", val="FollowedHyperlink")
        _sub(st, "uiPriority", val="99")
        rPr = _sub(st, "rPr")
        _sub(rPr, "color", val=BRAND_NAVY)
        _sub(rPr, "u", val="single")


def _add_lead_style(styles):
    """A custom 'BrandDocs Lead' intro-paragraph style (DOCX-A5): teal 13pt lead.

    Rounds out the editorial system with a styled intro paragraph the artifact
    catalog and style_details surface, alongside the untouched existing styles.
    """
    _add_paragraph_style(
        styles,
        "BrandDocsLead",
        "BrandDocs Lead",
        color=BRAND_TEAL,
        size_pt=13,
    )


def _add_table_compact_style(styles):
    """A SECOND custom ``w:type='table'`` style (DOCX-A7): BrandDocsTableCompact.

    Minimal navy hairline borders, no row banding, tighter cell margins, and a
    ``lastRow`` (totals) ``w:tblStylePr`` conditional format (teal fill + white
    bold). Authored ALONGSIDE BrandDocsTable (which stays ``table.default``); the
    resolver binds the first custom table style, so adding this second one after it
    is order-stable.
    """
    st = _sub(styles, "style", type="table", styleId="BrandDocsTableCompact")
    st.set(_w("customStyle"), "1")
    _sub(st, "name", val="BrandDocs Table Compact")
    _sub(st, "basedOn", val="TableNormal")
    _sub(st, "uiPriority", val="99")
    tblPr = _sub(st, "tblPr")
    borders = _sub(tblPr, "tblBorders")
    # Navy hairlines (no inside vertical rule -> a lighter, totals-style grid).
    for side in ("top", "bottom", "insideH"):
        _sub(borders, side, val="single", sz="2", space="0", color=BRAND_NAVY)
    margins = _sub(tblPr, "tblCellMar")
    _sub(margins, "top", w="40", type="dxa")
    _sub(margins, "bottom", w="40", type="dxa")
    _sub(margins, "left", w="80", type="dxa")
    _sub(margins, "right", w="80", type="dxa")
    rPr = _sub(st, "rPr")
    _sub(rPr, "color", val=BRAND_NAVY)
    # lastRow (totals) conditional format: teal fill + white bold text.
    lr = _sub(st, "tblStylePr", type="lastRow")
    lr_rpr = _sub(lr, "rPr")
    _sub(lr_rpr, "b")
    _sub(lr_rpr, "color", val=WHITE)
    lr_tcpr = _sub(lr, "tcPr")
    _sub(lr_tcpr, "shd", val="clear", color="auto", fill=BRAND_TEAL)
    return st


def _ensure_list_paragraph_style(styles):
    """python-docx's default styles.xml has no 'List Paragraph'; add a minimal one."""
    for st in styles.findall(_w("style")):
        if st.get(_w("styleId")) == "ListParagraph":
            return
    st = _sub(styles, "style", type="paragraph", styleId="ListParagraph")
    _sub(st, "name", val="List Paragraph")
    _sub(st, "basedOn", val="Normal")
    _sub(st, "uiPriority", val="34")
    _sub(st, "qFormat")
    pPr = _sub(st, "pPr")
    _sub(pPr, "ind", left="720")


def _ensure_toc_styles(styles):
    """Add TOC entry styles + TOCHeading + Caption + Footnote styles if absent."""
    have = {st.get(_w("styleId")) for st in styles.findall(_w("style"))}

    def add_simple(style_id, name, *, based_on="Normal", ui="39"):
        if style_id in have:
            return
        st = _sub(styles, "style", type="paragraph", styleId=style_id)
        _sub(st, "name", val=name)
        _sub(st, "basedOn", val=based_on)
        _sub(st, "uiPriority", val=ui)
        if style_id.startswith("TOC"):
            pPr = _sub(st, "pPr")
            _sub(pPr, "tabs")  # placeholder; real entries carry their own tabs

    for lvl in (1, 2, 3):
        add_simple(f"TOC{lvl}", f"TOC {lvl}")
    add_simple("TOCHeading", "TOC Heading", based_on="Heading1")
    add_simple("TableofFigures", "Table of Figures")
    if "Caption" not in have:
        st = _sub(styles, "style", type="paragraph", styleId="Caption")
        _sub(st, "name", val="Caption")
        _sub(st, "basedOn", val="Normal")
        _sub(st, "uiPriority", val="35")
        _sub(st, "qFormat")
        rPr = _sub(st, "rPr")
        _sub(rPr, "i")
        _sub(rPr, "color", val=BRAND_TEAL)
        _sub(rPr, "sz", val="18")
    if "FootnoteText" not in have:
        st = _sub(styles, "style", type="paragraph", styleId="FootnoteText")
        _sub(st, "name", val="Footnote Text")
        _sub(st, "basedOn", val="Normal")
        rPr = _sub(st, "rPr")
        _sub(rPr, "sz", val="20")
    if "FootnoteReference" not in have:
        st = _sub(styles, "style", type="character", styleId="FootnoteReference")
        _sub(st, "name", val="Footnote Reference")
        rPr = _sub(st, "rPr")
        _sub(rPr, "vertAlign", val="superscript")


def _build_styles(doc):
    styles = doc.styles.element
    # DOCX-A3: pin the brand body face/size/color as the document default run.
    _mutate_doc_defaults(doc)
    _ensure_list_paragraph_style(styles)
    _ensure_toc_styles(styles)
    # DOCX-A5: restyle the builtin Heading1/2/3 to the brand look + outlineLevel.
    _override_builtin_headings(styles)
    _add_section_number_style(styles)
    _add_lead_style(styles)
    # DOCX-A6: named character styles + a real Hyperlink/FollowedHyperlink style.
    _add_character_styles(styles)
    # Branded paragraph styles.
    _add_paragraph_style(
        styles,
        "BrandDocsCoverTitle",
        "BrandDocs Cover Title",
        color=BRAND_NAVY,
        bold=True,
        size_pt=28,
    )
    _add_paragraph_style(
        styles,
        "BrandDocsCoverSubtitle",
        "BrandDocs Cover Subtitle",
        color=BRAND_TEAL,
        size_pt=14,
    )
    _add_paragraph_style(
        styles,
        "BrandDocsCallout",
        "BrandDocs Callout",
        color=BRAND_NAVY,
        shading=BRAND_LIGHT,
        box_border=BRAND_TEAL,
        left_accent=BRAND_AMBER,
    )
    _add_paragraph_style(
        styles,
        "BrandDocsQuote",
        "BrandDocs Quote",
        color=BRAND_NAVY,
        shading=BRAND_BAND,
        box_border=BRAND_NAVY,
        left_accent=BRAND_AMBER,
        size_pt=12,
    )
    # List styles -> reference w:num 1 (bullet L1), 2 (bullet L2), 3 (number L1).
    # BUG-LIST-ILVL: the L2 bullet pins w:ilvl=1 so it binds to level 1 of
    # abstractNum 0 (a distinct ``list.bullet.2`` role, not a dedup of L1).
    _add_list_style(
        styles, "BrandDocsBulletL1", "BrandDocs Bullet L1", num_id=1, ilvl=0
    )
    _add_list_style(
        styles, "BrandDocsBulletL2", "BrandDocs Bullet L2", num_id=2, ilvl=1
    )
    _add_list_style(
        styles, "BrandDocsNumberL1", "BrandDocs Number L1", num_id=3, ilvl=0
    )
    # Branded table style (table.default) + a second compact/totals table style.
    _add_table_style(styles)
    _add_table_compact_style(styles)


# ---------------------------------------------------------------------------
# NUMBERING - a real ``word/numbering.xml`` with abstractNum + num.
# python-docx exposes no numbering authoring API for a fresh document, so the
# part is built with lxml and attached to the package.
# ---------------------------------------------------------------------------
def _populate_numbering(root) -> None:
    """Fill an existing ``w:numbering`` element with BrandDocs abstractNum/num defs.

    The python-docx default template already ships a (empty) ``word/numbering.xml``
    part, already related from ``document.xml``. We MUST reuse that part - adding a
    second part of the same name corrupts the zip - so this mutates the existing
    ``CT_Numbering`` element in place rather than authoring a fresh part.
    """
    for child in list(root):
        root.remove(child)

    def abstract(aid, levels):
        an = _sub(root, "abstractNum", abstractNumId=str(aid))
        _sub(an, "multiLevelType", val="hybridMultilevel")
        for lvl, level in enumerate(levels):
            fmt, text, indent = level[0], level[1], level[2]
            font = level[3] if len(level) > 3 else None
            lvl_el = _sub(an, "lvl", ilvl=str(lvl))
            _sub(lvl_el, "start", val="1")
            _sub(lvl_el, "numFmt", val=fmt)
            # BUG-LIST-BULLETGLYPH: an EXPLICIT, real Unicode lvlText glyph on a
            # plain text font. The old defs used Symbol-font private-use codepoints
            # (U+F0B7 / U+F0A7) which Word maps to a bullet/section mark but
            # LibreOffice renders as a stray club/box (no glyph at those PUA
            # points without the Symbol font). Plain Unicode bullets on a standard
            # font render identically in Word and LibreOffice.
            _sub(lvl_el, "lvlText", val=text)
            _sub(lvl_el, "lvlJc", val="left")
            pPr = _sub(lvl_el, "pPr")
            _sub(pPr, "ind", left=str(indent), hanging="360")
            if fmt == "bullet" and font:
                rPr = _sub(lvl_el, "rPr")
                rfonts = _sub(rPr, "rFonts")
                rfonts.set(_w("ascii"), font)
                rfonts.set(_w("hAnsi"), font)
                rfonts.set(_w("cs"), font)
                rfonts.set(_w("hint"), "default")
        return an

    # abstractNum 0: two-level bullet. L1 = filled round bullet (U+2022), L2 =
    # en-dash (U+2013), both on Arial so the glyphs are readable in Word AND
    # LibreOffice (no Symbol-font club/box).
    abstract(
        0,
        [
            ("bullet", "•", 720, "Arial"),
            ("bullet", "–", 1440, "Arial"),
        ],
    )
    # abstractNum 1: decimal numbered list.
    abstract(1, [("decimal", "%1.", 720)])

    # num bindings: 1->bullet(ilvl0 entry), 2->bullet, 3->decimal.
    for num_id, aid in ((1, 0), (2, 0), (3, 1)):
        n = _sub(root, "num", numId=str(num_id))
        _sub(n, "abstractNumId", val=str(aid))


# ---------------------------------------------------------------------------
# FOOTNOTES - a real ``word/footnotes.xml`` with the two reserved separators
# plus one authored footnote (id 2).
# ---------------------------------------------------------------------------
def _build_footnotes_xml() -> bytes:
    nsmap = {"w": W}
    root = etree.Element(_w("footnotes"), nsmap=nsmap)

    def sep(fid, kind):
        fn = _sub(root, "footnote", type=kind, id=str(fid))
        p = _sub(fn, "p")
        r = _sub(p, "r")
        _sub(r, kind if kind == "separator" else "continuationSeparator")

    s = _sub(root, "footnote", type="separator", id="-1")
    p = _sub(s, "p")
    r = _sub(p, "r")
    _sub(r, "separator")
    c = _sub(root, "footnote", type="continuationSeparator", id="0")
    p = _sub(c, "p")
    r = _sub(p, "r")
    _sub(r, "continuationSeparator")

    # The authored footnote (id 2).
    fn = _sub(root, "footnote", id="2")
    p = _sub(fn, "p")
    pPr = _sub(p, "pPr")
    _sub(pPr, "pStyle", val="FootnoteText")
    ref_run = _sub(p, "r")
    ref_rpr = _sub(ref_run, "rPr")
    _sub(ref_rpr, "rStyle", val="FootnoteReference")
    _sub(ref_run, "footnoteRef")
    txt = _sub(p, "r")
    t = _sub(txt, "t")
    t.text = " BrandDocs Corp is a fictional company used only for testing."
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _attach_part(doc, partname, content_type, xml_bytes, rel_type):
    """Register a new package part + content-type override + document rel.

    Returns the relationship id assigned to ``document.xml -> partname``.
    """
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    package = doc.part.package
    puri = PackURI("/" + partname)
    part = Part(puri, content_type, xml_bytes, package)
    rid = doc.part.relate_to(part, rel_type)
    return rid


# ---------------------------------------------------------------------------
# COVER - a block-level w:sdt title content control + placeholder paragraphs.
# ---------------------------------------------------------------------------
def _cover_title_sdt():
    """A block-level ``w:sdt`` cover-title content control (lxml; python-docx
    cannot author block-level SDTs).

    CR-COVER-TITLE: the SDT keeps ``w:alias='Title'`` (the strong cover-title
    anchor evidence ``cover._sdt_is_title`` keys on) but carries a REALISTIC
    synthetic title instead of the "Insert title here" prompt - the bare prompt
    read like an unfinished draft. The ``w:showingPlcHdr`` flag is dropped so Word
    treats the text as real content, not a greyed placeholder. Cover-anchor count
    is unaffected: ``discover_cover`` keys the SDT anchor on the container, not on
    the text, and the alias still classifies it as the title slot.
    """
    sdt = _el("sdt")
    sdtPr = _sub(sdt, "sdtPr")
    rpr = _sub(sdtPr, "rPr")
    _sub(rpr, "color", val=BRAND_NAVY)
    _sub(rpr, "sz", val="56")
    _sub(rpr, "b")
    _sub(sdtPr, "alias", val="Title")
    _sub(sdtPr, "tag", val="branddocs_title")
    _sub(sdtPr, "id", val="101")
    placeholder = _sub(sdtPr, "placeholder")
    _sub(placeholder, "docPart", val="DefaultPlaceholder_Title")
    _sub(sdtPr, "text")
    _sub(sdt, "sdtEndPr")
    sdtContent = _sub(sdt, "sdtContent")
    p = _sub(sdtContent, "p")
    pPr = _sub(p, "pPr")
    _sub(pPr, "pStyle", val="BrandDocsCoverTitle")
    r = _sub(p, "r")
    rpr2 = _sub(r, "rPr")
    _sub(rpr2, "color", val=BRAND_NAVY)
    _sub(rpr2, "sz", val="56")
    _sub(rpr2, "b")
    t = _sub(r, "t")
    t.text = "BrandDocs Brand Operations Review"
    return sdt


def _cover_background_shape():
    """A decorative DrawingML cover backdrop (DOCX-A9): an anchored navy rounded
    rectangle with a thin amber accent rule, returned as a ``w:drawing``.

    Authored as an ANCHORED ``wp:anchor`` ``wsp`` auto-shape (``a:prstGeom roundRect``,
    ``a:solidFill srgbClr navy``) that backs the navy COVER BAND region BELOW the
    date slot - NOT the SDT title. Hosting it on the band paragraph (a deliberately
    navy zone, well clear of the navy-on-white title) is what keeps the released
    cover look: the title stays navy-on-white and fully readable, with the navy
    banner below it. (BP-COVER-BACKDROP-FIX: the prior version anchored the shape at
    the top of the body behind the title, so the navy fill collided with the navy
    title text = dark-on-dark.) Decorative-only: its host paragraph carries no
    placeholder text so ``_paragraph_is_placeholder_slot`` never promotes it to a
    5th cover anchor; the shape still adds a real DrawingML auto-shape to the part
    walk without changing any role / anchor index. Static geometry + constant fill,
    so the bytes are reproducible (no raster, no wall-clock).
    """
    WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    drawing = _el("drawing")
    anchor = etree.SubElement(drawing, f"{{{WP}}}anchor")
    for k, v in (
        ("distT", "0"),
        ("distB", "0"),
        ("distL", "0"),
        ("distR", "0"),
        ("simplePos", "0"),
        ("relativeHeight", "0"),
        ("behindDoc", "1"),
        ("locked", "0"),
        ("layoutInCell", "1"),
        ("allowOverlap", "1"),
    ):
        anchor.set(k, v)
    sp = etree.SubElement(anchor, f"{{{WP}}}simplePos")
    sp.set("x", "0")
    sp.set("y", "0")
    hpos = etree.SubElement(anchor, f"{{{WP}}}positionH")
    hpos.set("relativeFrom", "column")
    etree.SubElement(hpos, f"{{{WP}}}posOffset").text = "0"
    vpos = etree.SubElement(anchor, f"{{{WP}}}positionV")
    vpos.set("relativeFrom", "paragraph")
    etree.SubElement(vpos, f"{{{WP}}}posOffset").text = "0"
    ext = etree.SubElement(anchor, f"{{{WP}}}extent")
    ext.set("cx", "5943600")  # ~6.5in usable width
    ext.set("cy", "274320")  # 0.3in band - sits over the band line, clears the table
    eff = etree.SubElement(anchor, f"{{{WP}}}effectExtent")
    for k in ("l", "t", "r", "b"):
        eff.set(k, "0")
    etree.SubElement(anchor, f"{{{WP}}}wrapNone")
    docpr = etree.SubElement(anchor, f"{{{WP}}}docPr")
    docpr.set("id", "400")
    docpr.set("name", "BrandDocsCoverBackdrop")
    etree.SubElement(anchor, f"{{{WP}}}cNvGraphicFramePr")
    graphic = etree.SubElement(anchor, f"{{{A}}}graphic")
    gdata = etree.SubElement(graphic, f"{{{A}}}graphicData")
    gdata.set("uri", WPS)
    wsp = etree.SubElement(gdata, f"{{{WPS}}}wsp")
    etree.SubElement(wsp, f"{{{WPS}}}cNvSpPr")
    sppr = etree.SubElement(wsp, f"{{{WPS}}}spPr")
    xfrm = etree.SubElement(sppr, f"{{{A}}}xfrm")
    off = etree.SubElement(xfrm, f"{{{A}}}off")
    off.set("x", "0")
    off.set("y", "0")
    extb = etree.SubElement(xfrm, f"{{{A}}}ext")
    extb.set("cx", "5943600")
    extb.set("cy", "274320")
    geom = etree.SubElement(sppr, f"{{{A}}}prstGeom")
    geom.set("prst", "roundRect")
    av = etree.SubElement(geom, f"{{{A}}}avLst")
    gd = etree.SubElement(av, f"{{{A}}}gd")
    gd.set("name", "adj")
    gd.set("fmla", "val 8000")
    fill = etree.SubElement(sppr, f"{{{A}}}solidFill")
    etree.SubElement(fill, f"{{{A}}}srgbClr").set("val", BRAND_NAVY)
    ln = etree.SubElement(sppr, f"{{{A}}}ln")
    ln.set("w", "19050")
    lnfill = etree.SubElement(ln, f"{{{A}}}solidFill")
    etree.SubElement(lnfill, f"{{{A}}}srgbClr").set("val", BRAND_AMBER)
    bodypr = etree.SubElement(wsp, f"{{{WPS}}}bodyPr")
    bodypr.set("rtlIns", "0")
    bodypr.set("bIns", "0")
    return drawing


def _cover_band(doc):
    """BP-COVER-BAND: a navy brand banner paragraph with an amber bottom rule.

    Carries ``w:pPr/w:shd`` navy fill + a thick amber ``w:pBdr`` bottom border, and
    HOSTS the decorative DOCX-A9 backdrop shape (a ``behindDoc`` navy rounded-rect
    with an amber border) so the auto-shape backs THIS navy band region - well below
    the navy-on-white title - rather than colliding with the title. Its text is
    EMPTY so ``_paragraph_is_placeholder_slot`` (which excludes empty paragraphs)
    never promotes it to a 5th cover anchor; placed after the date slot it leaves
    the ``sdt.0/para.1/para.2/para.3`` anchor indices stable.
    """
    p = doc.add_paragraph("")
    pPr = p._p.get_or_add_pPr()
    _sub(pPr, "shd", val="clear", color="auto", fill=BRAND_NAVY)
    pbdr = _sub(pPr, "pBdr")
    _sub(pbdr, "bottom", val="single", sz="24", space="2", color=BRAND_AMBER)
    _sub(pPr, "spacing", before="60", after="180")
    # DOCX-A9: the decorative backdrop auto-shape lives on the band run so it sits
    # behindDoc THIS navy band, never behind the title.
    run = p.add_run()
    run._r.append(_cover_background_shape())
    return p


def _build_cover(doc):
    body = doc.element.body
    # 1) Block-level SDT title (inserted as the very first body child).
    sdt = _cover_title_sdt()
    body.insert(0, sdt)

    # 2..n) Placeholder paragraphs for subtitle / doc-id / date, each in the
    # cover region (before the TOC), each a short single-line slot. The slot
    # demo values are realistic synthetic content (CR-COVER-VALUES); the SDT
    # title above now carries a realistic synthetic title (CR-COVER-TITLE) while
    # still classifying as the title slot via its ``w:alias='Title'``.
    sub_p = _p(
        doc,
        "Annual Brand Operations Review - BrandDocs Corp (synthetic)",
        "BrandDocsCoverSubtitle",
    )
    docid_p = doc.add_paragraph("Document ID: DSK-BR-2026-014")
    date_p = doc.add_paragraph("June 5, 2026")
    # A navy brand band sits BELOW the date slot (empty text, not an anchor). It
    # also HOSTS the DOCX-A9 decorative backdrop auto-shape (behindDoc) so the navy
    # rounded-rect backs THIS band region, never the title. The released cover look
    # is preserved: navy-on-white title ABOVE, navy banner BELOW.
    band_p = _cover_band(doc)
    # Move them right after the SDT (they were appended at the end of the body),
    # preserving order: SDT -> subtitle -> doc-id -> date -> band.
    for p in (band_p._p, date_p._p, docid_p._p, sub_p._p):
        body.remove(p)
        sdt.addnext(p)


def _cell_docproperty(cell, prop_name, cached):
    """Replace a table cell's text with a DOCPROPERTY complex field (DOCX-A10).

    Authors ``begin(dirty)/instrText ' DOCPROPERTY <name> \\* MERGEFORMAT '/separate
    /cached result/end`` directly into the cell's first paragraph, the same
    complex-field shape the TOC / SEQ / PAGE fields already use. The cached result
    keeps the cover rendering before Word recomputes the bound document property.
    """
    para = cell.paragraphs[0]
    pp = para._p
    # Drop the literal placeholder run python-docx put there when the cell was made.
    for r in list(pp.findall(_w("r"))):
        pp.remove(r)
    pp.append(_fldchar("begin", dirty=True))
    pp.append(_run(f" DOCPROPERTY {prop_name} \\* MERGEFORMAT ", instr=True))
    pp.append(_fldchar("separate"))
    pp.append(_run(cached))
    pp.append(_fldchar("end"))
    # The cached value reads as a brand navy figure (direct run color, DOCX-A4).
    _brand_runs(para, color=BRAND_NAVY)


def _attach_custom_properties(doc, props):
    """Author ``docProps/custom.xml`` carrying the cover KPI custom properties.

    A real custom-properties part bound to the package so the cover scorecard
    DOCPROPERTY fields resolve to live document properties. Static literal values
    (pinned, never ``now()``), so the bytes are deterministic. Surfaces in the
    part_catalog as ``docProps/custom.xml`` (a new field-family provenance part).
    """
    cp_ns = "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
    vt_ns = "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"
    fmtid = "{D5CDD505-2E9C-101B-9397-08002B2CF9AE}"
    root = etree.Element(f"{{{cp_ns}}}Properties", nsmap={None: cp_ns, "vt": vt_ns})
    for i, (name, value) in enumerate(props, start=2):
        prop = etree.SubElement(root, f"{{{cp_ns}}}property")
        prop.set("fmtid", fmtid)
        prop.set("pid", str(i))
        prop.set("name", name)
        lpwstr = etree.SubElement(prop, f"{{{vt_ns}}}lpwstr")
        lpwstr.text = value
    xml_bytes = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    _attach_part(
        doc,
        "docProps/custom.xml",
        "application/vnd.openxmlformats-officedocument.custom-properties+xml",
        xml_bytes,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties",
    )


def _build_cover_scorecard(doc):
    """A compact executive scorecard in the front matter.

    It is a table, not a paragraph, so it enriches the rendered cover without
    creating extra paragraph cover anchors. This gives the example a realistic
    corporate cover surface while keeping the cover-slot ids stable. The KPI VALUES
    are now driven by DOCPROPERTY complex fields bound to docProps/custom.xml
    (DOCX-A10), the way real corporate covers bind live values.
    """
    table = doc.add_table(rows=2, cols=3)
    table.style = "BrandDocs Table"
    _set_col_widths(table, (2160, 2160, 2160))
    labels = ("Brand health", "Template coverage", "Audit posture")
    for cell, label in zip(table.rows[0].cells, labels):
        cell.text = label
        _brand_runs(cell.paragraphs[0], font="Arial", color=WHITE)
    # Value row: the first two cells are DOCPROPERTY fields; the third is literal.
    value_cells = table.rows[1].cells
    _cell_docproperty(value_cells[0], "BrandHealth", "92")
    _cell_docproperty(value_cells[1], "FormatsAudited", "3 formats")
    value_cells[2].text = "Deep QA"
    _brand_runs(value_cells[2].paragraphs[0], color=BRAND_NAVY)
    doc.add_paragraph("")


def _build_cover_kpi_summary(doc):
    """A compact 'Cover KPI summary' 2-col totals table (DOCX-A7).

    Exercises the SECOND custom table style ``BrandDocsTableCompact`` and its
    ``lastRow`` conditional format: a small metric/value table whose final row is a
    totals row (teal fill + white bold via the style's ``tblStylePr lastRow``). The
    ``w:tblLook`` requests lastRow conditional formatting. Kept compact so the cover
    surface stays clean; both branded table styles are now exercised.
    """
    table = doc.add_table(rows=4, cols=2)
    table.style = "BrandDocs Table Compact"
    tblPr = table._tbl.tblPr
    _sub(
        tblPr,
        "tblLook",
        firstRow="0",
        lastRow="1",
        firstColumn="0",
        lastColumn="0",
        noHBand="1",
        noVBand="1",
    )
    rows = [
        ("Templates audited", "3"),
        ("Roles captured", "14"),
        ("Open QA findings", "0"),
        ("Total checks passed", "27"),
    ]
    for r, (label, value) in enumerate(rows):
        cells = table.rows[r].cells
        cells[0].text = label
        cells[1].text = value
        _brand_runs(cells[0].paragraphs[0], color=BRAND_NAVY)
        # The totals (last) row carries the conditional white-bold text.
        last = r == len(rows) - 1
        _brand_runs(cells[1].paragraphs[0], color=WHITE if last else BRAND_NAVY)
        if last:
            _brand_runs(cells[0].paragraphs[0], color=WHITE)
    _set_col_widths(table, (4320, 2160))
    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# INDEX FRONT MATTER - three real complex fields with cached demo entries.
# ---------------------------------------------------------------------------
def _toc_heading(doc, text):
    return _p(doc, text, "TOCHeading")


def _toc_entry(doc, label, page, *, style):
    """A cached TOC/index entry paragraph: a nested PAGEREF field + tab + page."""
    p = _p(doc, "", style)
    pp = p._p
    # The entry hyperlink text run.
    r = _sub(pp, "r")
    t = _sub(r, "t")
    t.text = label
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    # A right tab + a cached PAGEREF (begin / instr / separate / page / end).
    tab = _sub(pp, "r")
    _sub(tab, "tab")
    pp.append(_fldchar("begin"))
    pp.append(_run(f" PAGEREF _Toc{page:04d} \\h ", instr=True))
    pp.append(_fldchar("separate"))
    pp.append(_run(str(page)))
    pp.append(_fldchar("end"))
    return p


def _complex_field(doc, instr, cached_paragraph_builder):
    """Wrap a complex field around cached result paragraphs.

    Layout (Word's pattern for a multi-paragraph field):
      p0: [begin(dirty)] [instrText]
      p1..pn: cached entry paragraphs (the "result"), each a normal paragraph
      pE: [end]
    A ``separate`` fldChar precedes the cached result; the closing ``end`` sits
    in its own trailing paragraph.
    """
    # Field begin + instruction (its own paragraph).
    begin_p = doc.add_paragraph()
    bp = begin_p._p
    bp.append(_fldchar("begin", dirty=True))
    bp.append(_run(instr, instr=True))
    bp.append(_fldchar("separate"))
    # Cached result entries (real styled paragraphs in between).
    cached_paragraph_builder(doc)
    # Field end (its own paragraph).
    end_p = doc.add_paragraph()
    end_p._p.append(_fldchar("end"))


def _build_index_front_matter(doc):
    # --- Table of Contents (outline) ---
    _toc_heading(doc, "Table of Contents")

    def _toc_entries(d):
        _toc_entry(d, "1  Overview", 3, style="TOC1")
        _toc_entry(d, "1.1  Scope", 3, style="TOC2")
        _toc_entry(d, "1.2  Audience", 4, style="TOC2")
        _toc_entry(d, "2  Methodology", 5, style="TOC1")
        _toc_entry(d, "2.1  Data sources", 5, style="TOC2")
        _toc_entry(d, "3  Results", 6, style="TOC1")

    _complex_field(doc, 'TOC \\o "1-3" \\h \\z \\u ', _toc_entries)

    # --- Table of Tables ---
    _toc_heading(doc, "Table of Tables")

    def _tot_entries(d):
        # Cached entries mirror the two real SEQ Table captions in the body and
        # the landscape appendix (does NOT change the fields count).
        _toc_entry(
            d, "Table 1. BrandDocs FY2026 quarterly revenue", 5, style="TableofFigures"
        )
        _toc_entry(
            d, "Table 2. BrandDocs risk and readiness matrix", 6, style="TableofFigures"
        )
        _toc_entry(
            d, "Table 3. BrandDocs program rollout matrix", 7, style="TableofFigures"
        )

    _complex_field(doc, 'TOC \\h \\z \\c "Table" ', _tot_entries)

    # --- Table of Figures ---
    _toc_heading(doc, "Table of Figures")

    def _tof_entries(d):
        _toc_entry(d, "Figure 1. BrandDocs Corp wordmark", 2, style="TableofFigures")
        _toc_entry(d, "Figure 2. BrandDocs growth curve", 6, style="TableofFigures")

    _complex_field(doc, 'TOC \\h \\z \\c "Figure" ', _tof_entries)


# ---------------------------------------------------------------------------
# DEMO BODY - lists, table+caption, figure+caption, callout, footnote, demo
# heading content. Everything below the index front matter is the freeform body
# a generation would clear.
# ---------------------------------------------------------------------------
def _seq_caption(doc, prefix, seq_name, tail, *, style="Caption"):
    """A real ``SEQ`` caption paragraph: 'Prefix N. tail' with a live SEQ field."""
    p = _p(doc, "", style)
    pp = p._p
    pp.append(_run(f"{prefix} "))
    pp.append(_fldchar("begin"))
    pp.append(_run(f" SEQ {seq_name} \\* ARABIC ", instr=True))
    pp.append(_fldchar("separate"))
    pp.append(_run("1"))
    pp.append(_fldchar("end"))
    pp.append(_run(f". {tail}"))
    return p


def _build_lists(doc):
    # CR-LISTS: realistic two-level bullet list (two L1 each with two L2
    # children) + a numbered rollout sequence. DOCX-A4: heading runs carry direct
    # Arial + navy; list-item runs carry direct Calibri + navy at 11pt, so the
    # brand's real visible typography is captured per role and as body defaults.
    _pr(doc, "Brand operating principles", "Heading1", font="Arial", color=BRAND_NAVY)
    _pr(
        doc,
        "Consistency before customization",
        "BrandDocsBulletL1",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    _pr(
        doc,
        "One palette: navy, teal, amber",
        "BrandDocsBulletL2",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    _pr(
        doc,
        "Type scale fixed across all templates",
        "BrandDocsBulletL2",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    _pr(
        doc,
        "Templates are contracts, not suggestions",
        "BrandDocsBulletL1",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    _pr(
        doc,
        "Cover, contents, and indices are preserved",
        "BrandDocsBulletL2",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    _pr(
        doc,
        "Body content is regenerated per request",
        "BrandDocsBulletL2",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    _pr(doc, "Rollout sequence", "Heading2", font="Arial", color=BRAND_NAVY)
    _pr(
        doc,
        "Extract the template surface and brand profile",
        "BrandDocsNumberL1",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    _pr(
        doc,
        "Review the captured slots and index front matter",
        "BrandDocsNumberL1",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    _pr(
        doc,
        "Generate the branded document",
        "BrandDocsNumberL1",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )


def _build_table(doc):
    # CR-TABLE: header + four quarters of internally-consistent synthetic data.
    # BUG-TABLE-NUMFMT: one currency style ($X.XXM) and one percent style
    # (+X.X%) across every row, so the columns read as formatted figures.
    _pr(doc, "BrandDocs quarterly revenue", "Heading2", font="Arial", color=BRAND_NAVY)
    table = doc.add_table(rows=5, cols=4)
    table.style = "BrandDocs Table"
    # Tell Word which conditional formats to apply (first row + banding).
    tblPr = table._tbl.tblPr
    _sub(
        tblPr,
        "tblLook",
        firstRow="1",
        lastRow="0",
        firstColumn="0",
        lastColumn="0",
        noHBand="0",
        noVBand="1",
    )
    hdr = ("Quarter", "Revenue", "Growth", "Region")
    for c, label in zip(table.rows[0].cells, hdr):
        c.text = label
    data = [
        ("Q1 2026", "$3.20M", "+8.4%", "North"),
        ("Q2 2026", "$3.48M", "+8.8%", "South"),
        ("Q3 2026", "$3.91M", "+12.4%", "EMEA"),
        ("Q4 2026", "$4.25M", "+8.7%", "APAC"),
    ]
    for r, row in enumerate(data, start=1):
        for c, val in zip(table.rows[r].cells, row):
            c.text = val
    # DOCX-A4: brand the cell runs (Arial/white header, Calibri/navy body).
    _brand_table_body(table)
    _seq_caption(
        doc,
        "Table",
        "Table",
        "BrandDocs FY2026 quarterly revenue by region (synthetic data).",
    )


def _build_figure(doc, logo_rid, logo_cx, logo_cy):
    _p(doc, "BrandDocs wordmark", "Heading2")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run._r.append(
        _inline_drawing(logo_rid, logo_cx, logo_cy, "BrandDocsLogoFigure", 100)
    )
    _seq_caption(doc, "Figure", "Figure", "BrandDocs Corp wordmark (synthetic).")


def _build_callout(doc):
    # CR-CALLOUT: an on-brand synthetic note (not a meta tooling instruction).
    # DOCX-A4: the callout body run carries its brand navy directly; DOCX-A6: a
    # BrandDocsLeadIn small-caps navy lead-in opens it.
    p = _p(doc, "", "BrandDocsCallout")
    _add_styled_run(p, "Note. ", "BrandDocsLeadIn", color=BRAND_NAVY)
    run = p.add_run(
        "BrandDocs Corp is a synthetic, fictional company used to demonstrate an "
        "on-brand internal brief. All figures, names, and regions in this template "
        "are illustrative."
    )
    from docx.shared import RGBColor

    run.font.color.rgb = RGBColor.from_string(BRAND_NAVY)


def _build_editorial_quote(doc):
    # DOCX-A4: the quote body run carries brand navy directly; DOCX-A6: a
    # BrandDocsEmphasis teal-bold mark highlights the operative phrase.
    p = _p(doc, "", "BrandDocsQuote")
    from docx.shared import RGBColor

    run = p.add_run("Design rule: preserve proven structure only while it ")
    run.font.color.rgb = RGBColor.from_string(BRAND_NAVY)
    _add_styled_run(
        p, "improves the final artifact", "BrandDocsEmphasis", color=BRAND_TEAL
    )
    run2 = p.add_run(
        ". If visual QA shows blank space, stale content, or broken flow, "
        "repair the composition rather than copying the template blindly."
    )
    run2.font.color.rgb = RGBColor.from_string(BRAND_NAVY)


def _build_risk_matrix(doc):
    _pr(doc, "Readiness and risk matrix", "Heading2", font="Arial", color=BRAND_NAVY)
    _pr(
        doc,
        "The matrix below gives the template a realistic governance surface: it "
        "mixes status, owner, signal, and mitigation text in a branded table.",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    table = doc.add_table(rows=5, cols=4)
    table.style = "BrandDocs Table"
    tblPr = table._tbl.tblPr
    _sub(
        tblPr,
        "tblLook",
        firstRow="1",
        lastRow="0",
        firstColumn="0",
        lastColumn="0",
        noHBand="0",
        noVBand="1",
    )
    headers = ("Area", "Signal", "Owner", "Mitigation")
    for c, label in zip(table.rows[0].cells, headers):
        c.text = label
    rows = [
        ("Visual QA", "Renderer drift", "Platform", "Smoke-test DOCX/PPTX/XLSX"),
        (
            "Comprehension",
            "Missing inventory",
            "Model Ops",
            "Fail closed on empty refs",
        ),
        ("Brand fidelity", "Style fallback", "Design", "Re-assert captured styles"),
        ("Delivery", "Stale cache", "Engineering", "Refresh visible field results"),
    ]
    for r, row in enumerate(rows, start=1):
        for c, val in zip(table.rows[r].cells, row):
            c.text = val
    _brand_table_body(table)
    _set_col_widths(table, (1980, 1980, 1800, 3600))
    _seq_caption(
        doc, "Table", "Table", "BrandDocs risk and readiness matrix (synthetic)."
    )


def _build_footnote_paragraph(doc, footnote_id):
    _pr(doc, "BrandDocs footnote demo", "Heading2", font="Arial", color=BRAND_NAVY)
    body = doc.add_paragraph("BrandDocs Corp")
    body.add_run(" is a registered placeholder brand")
    # Append a footnoteReference run (id -> footnotes.xml).
    fr = _sub(body._p, "r")
    frpr = _sub(fr, "rPr")
    _sub(frpr, "rStyle", val="FootnoteReference")
    _sub(fr, "footnoteReference", id=str(footnote_id))
    body.add_run(" used throughout this template.")
    # DOCX-A4: the two text runs carry brand navy Calibri (the footnoteRef run has
    # no text, so _brand_runs skips it and the superscript style is preserved).
    _brand_runs(body, font="Calibri", color=BRAND_NAVY, size_hp=22)


def _build_demo_body(doc):
    # CR-DEMO-BODY: keep the H1 + para + H2 + para demo region, but replace the
    # lorem filler with readable synthetic prose. detect_demo_region keys on the
    # first body-region Heading-1 structurally (style id + captured own text),
    # so demo_region.present stays True with realistic copy. DOCX-A4 stamps direct
    # Arial/navy on the headings and Calibri/navy on the body runs; DOCX-A5 adds a
    # styled Lead intro paragraph below the H1.
    _pr(doc, "Overview", "Heading1", font="Arial", color=BRAND_NAVY)
    _pr(
        doc,
        "A complete, on-brand BrandDocs Corp template, captured and ready to drive "
        "branded generation.",
        "BrandDocsLead",
        font="Calibri",
        color=BRAND_TEAL,
        size_hp=26,
    )
    _pr(
        doc,
        "This brief summarizes how the BrandDocs Corp brand system is applied "
        "across internal documents. It exists to show a complete, on-brand "
        "template; a generation run replaces this body with the requested content.",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    _pr(doc, "Methodology", "Heading2", font="Arial", color=BRAND_NAVY)
    _pr(
        doc,
        "Figures are drawn from a synthetic operations dataset maintained by the "
        "BrandDocs brand office. Revenue, growth, and regional splits are "
        "illustrative and should not be read as real performance data.",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    # A real external hyperlink (w:hyperlink -> external relationship): a near-
    # universal template element, and one the generator now authors natively too.
    link_p = doc.add_paragraph("Full brand guidelines: ")
    _brand_runs(link_p, font="Calibri", color=BRAND_NAVY, size_hp=22)
    # The link run itself keeps the Hyperlink rStyle (teal) authored in
    # _add_external_hyperlink (DOCX-A6); only the lead-in run is branded here.
    _add_external_hyperlink(
        doc, link_p, "https://example.com/brand", "brand.example.com"
    )


# ---------------------------------------------------------------------------
# DRAWINGS - an inline picture (figure) and a header logo, both referencing the
# same image part by relationship id.
# ---------------------------------------------------------------------------
def _inline_drawing(rid, cx, cy, name, doc_pr_id):
    drawing = _el("drawing")
    inline = etree.SubElement(drawing, f"{{{WP}}}inline")
    inline.set("distT", "0")
    inline.set("distB", "0")
    inline.set("distL", "0")
    inline.set("distR", "0")
    ext = etree.SubElement(inline, f"{{{WP}}}extent")
    ext.set("cx", str(cx))
    ext.set("cy", str(cy))
    eff = etree.SubElement(inline, f"{{{WP}}}effectExtent")
    for k in ("l", "t", "r", "b"):
        eff.set(k, "0")
    docpr = etree.SubElement(inline, f"{{{WP}}}docPr")
    docpr.set("id", str(doc_pr_id))
    docpr.set("name", name)
    cnv = etree.SubElement(inline, f"{{{WP}}}cNvGraphicFramePr")
    locks = etree.SubElement(cnv, f"{{{A}}}graphicFrameLocks")
    locks.set("noChangeAspect", "1")
    graphic = etree.SubElement(inline, f"{{{A}}}graphic")
    gdata = etree.SubElement(graphic, f"{{{A}}}graphicData")
    gdata.set("uri", PIC)
    pic = etree.SubElement(gdata, f"{{{PIC}}}pic")
    nvpic = etree.SubElement(pic, f"{{{PIC}}}nvPicPr")
    cnvpr = etree.SubElement(nvpic, f"{{{PIC}}}cNvPr")
    cnvpr.set("id", "0")
    cnvpr.set("name", name)
    etree.SubElement(nvpic, f"{{{PIC}}}cNvPicPr")
    blipfill = etree.SubElement(pic, f"{{{PIC}}}blipFill")
    blip = etree.SubElement(blipfill, f"{{{A}}}blip")
    blip.set(f"{{{R}}}embed", rid)
    stretch = etree.SubElement(blipfill, f"{{{A}}}stretch")
    etree.SubElement(stretch, f"{{{A}}}fillRect")
    sppr = etree.SubElement(pic, f"{{{PIC}}}spPr")
    xfrm = etree.SubElement(sppr, f"{{{A}}}xfrm")
    off = etree.SubElement(xfrm, f"{{{A}}}off")
    off.set("x", "0")
    off.set("y", "0")
    extb = etree.SubElement(xfrm, f"{{{A}}}ext")
    extb.set("cx", str(cx))
    extb.set("cy", str(cy))
    geom = etree.SubElement(sppr, f"{{{A}}}prstGeom")
    geom.set("prst", "rect")
    etree.SubElement(geom, f"{{{A}}}avLst")
    return drawing


def _build_header_footer(doc, logo_rid, cx, cy):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    run = hp.add_run()
    # The header logo references the header part's OWN relationship to the image.
    run._r.append(_inline_drawing(logo_rid, cx, cy, "BrandDocsHeaderLogo", 200))

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = "Page "
    fpp = fp._p
    fpp.append(_fldchar("begin"))
    fpp.append(_run(" PAGE ", instr=True))
    fpp.append(_fldchar("separate"))
    fpp.append(_run("1"))
    fpp.append(_fldchar("end"))


def _relate_image_to(part, image_blob, partname="media/image1.png"):
    """Ensure ``part`` (document or header) relates to a shared image part.

    The image part is added once to the package; each consuming part gets its own
    relationship id back.
    """
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    package = part.package
    puri = PackURI("/word/" + partname)
    image_part = None
    for p in package.iter_parts():
        if p.partname == puri:
            image_part = p
            break
    if image_part is None:
        image_part = Part(puri, "image/png", image_blob, package)
    rid = part.relate_to(image_part, f"{R}/image")
    return rid


# ---------------------------------------------------------------------------
# SECTIONS - convert the document into a portrait section followed by a
# landscape section.
# ---------------------------------------------------------------------------
def _set_col_widths(table, widths_twips):
    """Pin explicit per-column widths (Twips) on a table so a wide landscape
    table fits the usable page width without column clipping (PL-PRINT-MARGINS).

    Disables autofit and stamps ``w:tcW`` (dxa) on every cell of every column,
    which is how Word honours fixed column widths.
    """
    table.autofit = False
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    layout = _sub(tblPr, "tblLayout")
    layout.set(_w("type"), "fixed")
    for row in table.rows:
        for cell, w_tw in zip(row.cells, widths_twips):
            cell.width = Twips(w_tw)


def _add_two_column_section(doc):
    """A THIRD section: a 2-column (newspaper) synopsis (DOCX-A8).

    Inserts a ``w:sectPr/w:cols w:num='2' w:space='720' w:sep='1'`` section break
    after the portrait body, carrying a short 'BrandDocs in brief' two-column
    synopsis. python-docx ``add_section`` creates the break; the ``w:cols`` child is
    stamped via lxml (python-docx has no column API). Adds a third ``w:sectPr`` with
    column geometry to the section inventory; the body skeleton stays valid.
    """
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    # Portrait geometry (the landscape section that follows sets its own size).
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Twips(12240)
    section.page_height = Twips(15840)
    sectPr = section._sectPr
    # Replace any existing w:cols with a 2-column spec (num/space/sep).
    existing = sectPr.find(_w("cols"))
    if existing is not None:
        sectPr.remove(existing)
    cols = _sub(sectPr, "cols", num="2", space="720", sep="1")
    cols.set(_w("equalWidth"), "1")
    _p(doc, "BrandDocs in brief", "Heading2")
    _pr(
        doc,
        "BrandDocs Corp is a synthetic brand used to exercise the on-brand "
        "document engine end to end. The template carries a cover, a contents "
        "block, indexed front matter, branded tables, figures, callouts, and a "
        "landscape appendix.",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    _pr(
        doc,
        "Everything visible here is illustrative. A generation run preserves the "
        "proven structure and replaces the body with the requested content while "
        "keeping the captured brand styles, palette, and typography intact.",
        font="Calibri",
        color=BRAND_NAVY,
        size_hp=22,
    )
    return section


def _add_titlepg_first_page_header(doc, logo_blob):
    """Give the FIRST section a distinct first-page header (DOCX-A8).

    Sets ``w:titlePg`` on the first section and authors a first-page header part
    (a brand band paragraph) wired via a ``first`` headerReference. python-docx
    exposes ``section.first_page_header`` (creating the part + reference + rels), so
    the new ``header*.xml`` part lands in the part_catalog. Decorative brand band
    only (an amber bottom rule), so it never changes a role or anchor.
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    fph = section.first_page_header
    fph.is_linked_to_previous = False
    para = fph.paragraphs[0]
    para.text = "BrandDocs Corp"
    _brand_runs(para, font="Arial", color=BRAND_NAVY)
    pPr = para._p.get_or_add_pPr()
    pbdr = _sub(pPr, "pBdr")
    _sub(pbdr, "bottom", val="single", sz="18", space="2", color=BRAND_AMBER)


def _add_landscape_section(doc, curve_rid=None, fig_cx=0, fig_cy=0):
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    # Swap page width/height for landscape (python-docx does not auto-swap).
    new_section.page_width = Twips(15840)  # 11"
    new_section.page_height = Twips(12240)  # 8.5"
    _p(doc, "BrandDocs landscape appendix", "Heading1")
    doc.add_paragraph(
        "This appendix sits in a landscape section. It carries a wide program "
        "rollout matrix and a second brand figure that the portrait body cannot "
        "fit. Demo content the generator may clear."
    )

    # PL-LANDSCAPE-CONTENT: a wide 7-column synthetic "program rollout matrix",
    # reusing the same custom table style + tblLook (banding exercised again).
    _p(doc, "Program rollout matrix", "Heading2")
    wide = doc.add_table(rows=5, cols=7)
    wide.style = "BrandDocs Table"
    wtblPr = wide._tbl.tblPr
    _sub(
        wtblPr,
        "tblLook",
        firstRow="1",
        lastRow="0",
        firstColumn="0",
        lastColumn="0",
        noHBand="0",
        noVBand="1",
    )
    whdr = ("Workstream", "Owner", "Q1", "Q2", "Q3", "Q4", "Status")
    for c, label in zip(wide.rows[0].cells, whdr):
        c.text = label
    wdata = [
        (
            "Template surface",
            "Brand Office",
            "Scope",
            "Build",
            "QA",
            "Ship",
            "On track",
        ),
        (
            "Profile extraction",
            "Platform",
            "Spec",
            "Build",
            "Verify",
            "Tune",
            "On track",
        ),
        (
            "Generation engine",
            "Platform",
            "Design",
            "Build",
            "Build",
            "Verify",
            "At risk",
        ),
        (
            "Rollout & training",
            "Enablement",
            "Plan",
            "Draft",
            "Pilot",
            "Launch",
            "Planned",
        ),
    ]
    for r, row in enumerate(wdata, start=1):
        for c, val in zip(wide.rows[r].cells, row):
            c.text = val
    # PL-PRINT-MARGINS: fixed column widths sum to 13680 twips (15840 - 2*1in
    # margins), so the 7-column table fits the landscape usable width.
    _set_col_widths(wide, (2880, 2400, 1680, 1680, 1680, 1680, 1680))
    _seq_caption(
        doc, "Table", "Table", "BrandDocs FY2026 program rollout matrix (synthetic)."
    )

    # PL-LANDSCAPE-CAPTION-SEQ / CR-FIG2-CURVE: a real SECOND inline figure so the
    # cached "Figure 2" in the Table of Figures corresponds to a rendered figure.
    # This is a DISTINCT media part (word/media/image2.png, the growth-curve PNG),
    # NOT a reuse of the logo - Figure 2 is a real rising growth curve, not the
    # wordmark. The curve PNG is ~12:5, so the inline extent matches that aspect.
    if curve_rid is not None:
        _p(doc, "BrandDocs growth curve", "Heading2")
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run()
        frun._r.append(
            _inline_drawing(curve_rid, fig_cx, fig_cy, "BrandDocsGrowthFigure", 300)
        )
        _seq_caption(
            doc,
            "Figure",
            "Figure",
            "BrandDocs Corp growth curve (synthetic illustration).",
        )
    return new_section


# ---------------------------------------------------------------------------
# settings.xml - footnotePr + a docId-free, deterministic settings part already
# exists; we just request field update on open so cached TOC/SEQ recompute.
# ---------------------------------------------------------------------------
def _request_update_fields(doc):
    settings = doc.settings.element
    if settings.find(_w("updateFields")) is None:
        uf = _el("updateFields", val="true")
        settings.insert(0, uf)


def build(out: Path = OUT) -> Path:
    doc = Document()  # python-docx default template (single portrait section)

    # 1) Styles (paragraph + list + table) authored into styles.xml; docDefaults
    # body run pinned to the brand (DOCX-A3 inside _build_styles).
    _build_styles(doc)

    # 1b) DOCX-A1 + DOCX-A2: rewrite the shipped word/theme/theme1.xml in place so
    # the fontScheme (Arial major / Calibri minor) and clrScheme (BrandDocs
    # navy/teal/amber/light) carry the brand at the theme level.
    _mutate_theme(doc)

    # 2) Numbering: the default template already ships an (empty) numbering part,
    # already related from document.xml. Populate it in place (a duplicate part
    # of the same name would corrupt the zip).
    _populate_numbering(doc.part.numbering_part.element)

    # 3) Footnotes part attached; footnote id 2 is the authored note.
    _attach_part(
        doc,
        "word/footnotes.xml",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
        _build_footnotes_xml(),
        f"{R}/footnotes",
    )

    # 4) Shared logo image: a generated text-only BrandDocs wordmark. The SAME
    # wordmark every BrandDocs template embeds (CR-LOGO-MARK). Related to the
    # document part (for Figure 1) and to the header part (for the header logo)
    # independently. A SECOND media part carries the real growth-curve figure
    # (CR-FIG2-CURVE).
    logo_blob = branddocs_mark_png(640, 160)  # wide wordmark (image1.png)
    curve_blob = branddocs_curve_png(480, 200)  # ~12:5 growth curve (image2.png)
    doc_logo_rid = _relate_image_to(doc.part, logo_blob)
    doc_curve_rid = _relate_image_to(doc.part, curve_blob, partname="media/image2.png")

    # 4b) DOCX-A10: custom document properties the cover scorecard DOCPROPERTY
    # fields bind to (pinned literal values, deterministic).
    _attach_custom_properties(
        doc, (("BrandHealth", "92"), ("FormatsAudited", "3 formats"))
    )

    # 5) Cover (block-level SDT title + placeholder slots), then the three index
    # fields, then the demo body. add_* appends in document order, so build the
    # front matter first, then the body. The scorecard KPI values are DOCPROPERTY
    # fields (DOCX-A10); a compact KPI summary exercises the second table style
    # (DOCX-A7).
    _build_cover(doc)
    _build_cover_scorecard(doc)
    _build_cover_kpi_summary(doc)
    _build_index_front_matter(doc)

    # Body content (freeform; a generation clears this region).
    _build_lists(doc)
    _build_table(doc)
    # Figure 1 (the wordmark) uses the document-part image relationship. Preserve
    # its 4:1 aspect ratio so the text never stretches.
    fig_cx, fig_cy = 2743200, 685800
    _build_figure(doc, doc_logo_rid, fig_cx, fig_cy)
    _build_callout(doc)
    _build_editorial_quote(doc)
    _build_risk_matrix(doc)
    _build_footnote_paragraph(doc, footnote_id=2)
    _build_demo_body(doc)

    # 5b) DOCX-A8: a 2-column (newspaper) synopsis section between the portrait
    # body and the landscape appendix - a third w:sectPr with column geometry.
    _add_two_column_section(doc)

    # 6) Header (logo) + footer (PAGE field). The header part needs its OWN
    # relationship to the shared image part.
    section = doc.sections[0]
    header_logo_rid = _relate_image_to(section.header.part, logo_blob)
    # Header logo is the SAME 4:1 wordmark, kept compact and undistorted.
    _build_header_footer(doc, header_logo_rid, 914400, 228600)
    # DOCX-A8: a distinct first-page header on the first section (w:titlePg + a
    # `first` headerReference), adding a header part to the part_catalog.
    _add_titlepg_first_page_header(doc, logo_blob)

    # 7) A second, landscape section after the portrait body. It carries a wide
    # rollout-matrix table and a real second figure (the cached "Figure 2"),
    # which is the DISTINCT growth-curve media part (word/media/image2.png).
    # The curve PNG is 480x200 (~12:5), so the extent matches: cx=2286000 EMU
    # (2.5in), cy=2286000*5//12=952500 EMU (exact 12:5, no distortion).
    land_fig_cx, land_fig_cy = 2286000, 952500
    _add_landscape_section(
        doc, curve_rid=doc_curve_rid, fig_cx=land_fig_cx, fig_cy=land_fig_cy
    )

    # 8) Ask Word to refresh the cached fields on open.
    _request_update_fields(doc)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    freeze_ooxml(out)
    return out


if __name__ == "__main__":
    path = build()
    print(f"built {path}")
